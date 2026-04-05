import logging
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .chart_mapper import ChartMapper

logger = logging.getLogger(__name__)

class DOCXExporter:
    """
    The 'Narrative Engine' for Professional Documentation.
    Focuses on text-first hierarchy and detailed formatting.
    """
    def __init__(self):
        self.doc = Document()
        self.ACCENT_COLOR = RGBColor(0x8B, 0x5C, 0xF6) # Violet 500

    def _set_header(self, text, level=1):
        h = self.doc.add_heading(text, level=level)
        run = h.runs[0]
        run.font.color.rgb = self.ACCENT_COLOR if level <= 1 else RGBColor(0,0,0)

    def generate_report(self, context: dict) -> io.BytesIO:
        """Orchestrates Word document generation."""
        # 1. Title
        self._set_header(context['session_title'], level=0)
        self.doc.add_paragraph(f"Report Date: {context['date']}")
        self.doc.add_paragraph(f"Dataset: {context['dataset_info']['name']}")
        
        # 2. Executive Summary
        self._set_header("Executive Summary", level=1)
        for bullet in context['executive_summary']:
            p = self.doc.add_paragraph(bullet, style='List Bullet')
            p.paragraph_format.space_after = Pt(6)

        # 3. Insights
        self._set_header("Analytical Insights", level=1)
        for insight in context['insights']:
            self._set_header(insight['title'], level=2)
            
            # Embed Visual (Chart or Table)
            visual_placed = False
            for artifact in insight.get('artifacts', []):
                # Word prefers images (static) for documents
                img_stream = ChartMapper.generate_static_image(artifact)
                if img_stream:
                    self.doc.add_picture(img_stream, width=Inches(5.5))
                    visual_placed = True
                    break
            
            # Key Takeaways
            self.doc.add_paragraph("Key Takeaways:", style='Body Text').bold = True
            for takeaway in insight.get('takeaways', []):
                self.doc.add_paragraph(takeaway, style='List Bullet')

        # 4. Appendix (Tables)
        self._set_header("Appendix: Detailed Data", level=1)
        for insight in context['insights']:
            data = insight.get('data', {})
            if isinstance(data, dict) and len(data) > 0:
                self._set_header(f"Data Table: {insight['tool']}", level=3)
                # Create simple table (truncate at 10 rows for professionalism)
                table = self.doc.add_table(rows=1, cols=min(len(data.keys()), 5))
                table.style = 'Table Grid'
                # Headers
                hdr_cells = table.rows[0].cells
                for i, key in enumerate(list(data.keys())[:5]):
                    hdr_cells[i].text = str(key)
                
        # Save to stream
        docx_stream = io.BytesIO()
        self.doc.save(docx_stream)
        docx_stream.seek(0)
        return docx_stream
